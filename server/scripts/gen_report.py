import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import date

doc = Document()

# 页面设置
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# 样式
style = doc.styles['Normal']
font = style.font
font.name = 'SimSun'
font.size = Pt(11)

# ===================== 封面 =====================
doc.add_paragraph('\n\n\n')
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('OpsMind 运维数字员工系统\n问题修复与优化报告')
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f'日期：{date.today().strftime("%Y年%m月%d日")}\n').font.size = Pt(12)
meta.add_run('环境：Windows 11 / Go 1.22+ / Vue 3.4+ / PostgreSQL 18 + pgvector\n').font.size = Pt(12)
meta.add_run('LLM：DeepSeek (deepseek-chat) / Embedding：硅基流动 (BAAI/bge-m3)').font.size = Pt(12)

doc.add_page_break()

# ===================== 目录页 =====================
doc.add_heading('目录', level=1)
toc_items = [
    '1. 申告状态机修复（reopen / resolve）',
    '2. 审计日志补全（Ticket / User / Role）',
    '3. LLM 与 Embedding 配置修复',
    '4. 知识库文章列表不显示修复',
    '5. 知识库种子数据填充',
    '6. 账号密码重置汇总',
    '7. 修改文件清单',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ===================== 正文 =====================

# 1. 申告状态机修复
doc.add_heading('1. 申告状态机修复', level=1)

doc.add_heading('1.1 问题描述', level=2)
doc.add_paragraph(
    '后台申告处理页面（TicketDetail.vue）存在两个状态转换缺陷：\n'
    '(1) 缺少"重新处理"（reopen）操作，已解决(4)和已关闭(5)的工单无法重新打开；\n'
    '(2) "已解决"（resolve）操作仅允许从处理中(2)状态执行，但前端在需补充信息(3)状态也显示了已解决按钮，导致点击后返回 400 错误。'
)

doc.add_heading('1.2 修复内容', level=2)
doc.add_paragraph(
    '后端 ticket_service.go：\n'
    '  • 新增 case "reopen"：允许状态 4/5 → 2，操作类型记录为 reopen\n'
    '  • 修改 case "resolve"：允许状态 2 或 3 → 4\n\n'
    '前端 TicketDetail.vue：\n'
    '  • 新增 reopen 按钮及对应 CSS 样式\n'
    '  • 状态 4/5 时显示"重新处理"按钮\n'
    '  • 补充 .record-action.reopen / .close / .auto_close 样式\n\n'
    '前端 utils/ticket.ts：\n'
    '  • actionText 映射新增 reopen 和 auto_close 条目'
)

doc.add_heading('1.3 申告状态机（完整）', level=2)
table = doc.add_table(rows=6, cols=3, style='Light Shading Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = '操作'
hdr[1].text = '允许的前置状态'
hdr[2].text = '目标状态'
data = [
    ('start（接单）', '待处理(1)', '处理中(2)'),
    ('request_info（需补充）', '处理中(2)', '需补充信息(3)'),
    ('resolve（已解决）', '处理中(2) / 需补充信息(3)', '已解决(4)'),
    ('reopen（重新处理）', '已解决(4) / 已关闭(5)', '处理中(2)'),
    ('close（关闭）', '任意状态', '已关闭(5)'),
]
for i, (action, from_s, to_s) in enumerate(data):
    row = table.rows[i+1].cells
    row[0].text = action
    row[1].text = from_s
    row[2].text = to_s

doc.add_page_break()

# 2. 审计日志
doc.add_heading('2. 审计日志补全', level=1)

doc.add_heading('2.1 问题描述', level=2)
doc.add_paragraph(
    '用户管理、角色管理、申告管理的关键操作均缺少审计日志写入，不符合安全合规要求。'
    '操作人身份、操作时间、操作内容无法追溯。'
)

doc.add_heading('2.2 修复内容', level=2)
doc.add_paragraph(
    'TicketService：\n'
    '  • 新增 auditRepo 字段和 writeAudit() 私有方法\n'
    '  • 所有状态变更、补充信息、添加记录、自动关闭均写入审计日志\n'
    '  • 目标类型：ticket，Action 格式：ticket:start / ticket:resolve 等\n\n'
    'UserService：\n'
    '  • 新增 auditRepo 字段和 writeAudit() 私有方法\n'
    '  • Create / Update / Freeze / Restore 均写入审计日志\n'
    '  • 方法签名新增 operatorID int64 参数\n\n'
    'RoleService：\n'
    '  • 新增 auditRepo 字段和 writeAudit() 私有方法\n'
    '  • Create / Update / Delete 均写入审计日志\n\n'
    'Handler 层适配：\n'
    '  • user.go / role.go 各 Handler 方法从 JWT 提取 operatorID 传给 Service\n\n'
    'main.go 初始化：\n'
    '  • 为 UserService、RoleService、TicketService 注入 auditRepo\n\n'
    '关键修复：writeAudit 中的 detail 必须使用 json.Marshal 包装为合法 JSON，'
    '不能直接传原始字符串给 datatypes.JSON，否则 PostgreSQL JSONB 列写入失败。'
)

doc.add_heading('2.3 审计日志覆盖范围', level=2)
table2 = doc.add_table(rows=12, cols=3, style='Light Shading Accent 1')
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr2 = table2.rows[0].cells
hdr2[0].text = '模块'
hdr2[1].text = '操作'
hdr2[2].text = '目标类型'
data2 = [
    ('申告', 'start（接单）', 'ticket'),
    ('申告', 'resolve（解决）', 'ticket'),
    ('申告', 'request_info（请求补充）', 'ticket'),
    ('申告', 'reopen（重新处理）', 'ticket'),
    ('申告', 'close（关闭）', 'ticket'),
    ('申告', 'auto_close（自动关闭）', 'ticket'),
    ('申告', 'supplement（补充信息）', 'ticket'),
    ('用户', 'create / update / freeze / restore', 'user'),
    ('角色', 'create / update / delete', 'role'),
    ('申告', 'AddRecord（处理记录）', 'ticket'),
]
for i, (mod, action, target) in enumerate(data2):
    row = table2.rows[i+1].cells
    row[0].text = mod
    row[1].text = action
    row[2].text = target

doc.add_page_break()

# 3. LLM & Embedding 配置修复
doc.add_heading('3. LLM 与 Embedding 配置修复', level=1)

doc.add_heading('3.1 问题描述', level=2)
doc.add_paragraph(
    '用户在门户端智能问答页面提问时，大模型卡住无响应。排查发现两个问题叠加：\n\n'
    '(1) LLM 配置错配：.env 文件配置了 DeepSeek API 作为 LLM 提供商，但数据库 llm_configs 表的默认配置指向 '
    'llama-cpp 本地服务（http://llama-cpp:8080/v1, qwen3-4b），而 llama-cpp 容器并未启动。'
    'LLMConfigManager 启动时从 DB 加载配置覆盖了 .env 的模型名，导致实际向 DeepSeek 发请求时使用了不存在的模型名。\n\n'
    '(2) 向量嵌入缺失：知识库 5 个分块的 embedding 列全为 NULL。文章发布时 Embedding API 调用失败但未报错，'
    '导致 pgvector 向量检索返回 NULL 分数，RAG 管道崩溃。'
)

doc.add_heading('3.2 修复内容', level=2)
doc.add_paragraph(
    'LLM 配置：\n'
    '  • 将 llm_configs 表中默认配置的 base_url 改为 https://api.deepseek.com/v1\n'
    '  • 模型名改为 deepseek-chat，API Key 填入真实密钥\n'
    '  • 重启服务器加载新配置\n\n'
    '向量嵌入：\n'
    '  • 编写 gen_embeddings.py 脚本，调用硅基流动 BAAI/bge-m3 API\n'
    '  • 为 5 个已有分块生成 1024 维向量，以 halfvec 格式写入 knowledge_chunks 表\n'
    '  • 验证向量检索恢复正常'
)

doc.add_heading('3.3 修复后 RAG 管道验证', level=2)
doc.add_paragraph(
    '提问"VPN密码怎么重置"，管道执行日志：\n'
    '  查询改写 → 多路检索 → 向量检索 → BM25检索 → 混合融合 → LLM生成 → 流式输出\n'
    'DeepSeek 正常返回 token 级流式回答，管道全部 6 步骤成功执行。'
)

doc.add_page_break()

# 4. 知识库文章列表不显示
doc.add_heading('4. 知识库文章列表不显示修复', level=1)

doc.add_heading('4.1 问题描述', level=2)
doc.add_paragraph(
    '后台管理端 /admin/knowledge 页面，进入 IT 运维 FAQ 知识库后文章列表显示为空，'
    '但数据库中实际存在 12 篇文章。'
)

doc.add_heading('4.2 根因分析', level=2)
doc.add_paragraph(
    '文件：server/internal/handler/knowledge.go:286\n\n'
    '原代码：status, _ := strconv.Atoi(c.DefaultQuery("status", "0"))\n\n'
    '前端在"全部状态"选项时不传 status 参数，后端使用默认值 "0"。'
    'Repository 层判断 status >= 0 时添加 WHERE status = 0 过滤条件。'
    '但文章状态枚举从 1 开始（1=草稿, 2=待审核, 3=已通过, 4=已发布），'
    'status=0 没有匹配任何文章，导致返回空列表。\n\n'
    '修复：将默认值从 "0" 改为 "-1"，-1 在 Repo 层表示不过滤。'
)

doc.add_heading('4.3 修复效果', level=2)
doc.add_paragraph('修复后 /admin/knowledge 正常显示全部 12 篇文章（6 篇原有 + 6 篇新增）。')

doc.add_page_break()

# 5. 知识库种子数据
doc.add_heading('5. 知识库种子数据填充', level=1)

doc.add_heading('5.1 背景', level=2)
doc.add_paragraph(
    '原知识库仅有 3 篇有效的 IT 运维 FAQ 文章（VPN 重置、WiFi 连接、Outlook 邮件），'
    '内容覆盖范围有限，无法支撑运维场景智能问答的实际需求。'
    '用户提问"服务器 A01 登录超时"时无法命中相关文章。'
)

doc.add_heading('5.2 新增文章', level=2)
articles = [
    ('服务器 SSH 连接超时排查指南', '排查 SSH 超时的 5 大原因：网络连通性、服务状态、DNS、负载、连接数限制'),
    ('MySQL 数据库连接超时故障处理', 'MySQL 连接超时的 6 步排查：服务状态、连接数、超时配置、网络、慢查询、连接池'),
    ('Nginx 502 Bad Gateway 排查流程', 'Nginx 502 错误的 5 步排查：上游服务、错误日志、超时配置、性能、应急措施'),
    ('Linux 磁盘空间不足应急处理', '磁盘满的应急处理 5 步：定位大文件、清理内容、释放 deleted 文件、扩容、预防'),
    ('Docker 容器常见故障排查', '容器故障 5 大类：无法启动、频繁重启、网络不通、磁盘满、资源限制'),
    ('常见网络故障排查方法', '网络故障系统排查 6 步：连通性、DNS、端口服务、防火墙、抓包、性能测试'),
]

table3 = doc.add_table(rows=len(articles)+1, cols=2, style='Light Shading Accent 1')
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
table3.rows[0].cells[0].text = '文章标题'
table3.rows[0].cells[1].text = '内容概要'
for i, (title, summary) in enumerate(articles):
    table3.rows[i+1].cells[0].text = title
    table3.rows[i+1].cells[1].text = summary

doc.add_heading('5.3 技术细节', level=2)
doc.add_paragraph(
    '使用 Go 编写的种子数据工具（cmd/seed_knowledge/main.go），通过 Embedding API 自动完成：\n'
    '  1. 创建文章（status=1 草稿）\n'
    '  2. 审核通过（status 1→2）\n'
    '  3. 分块（chunker.Split）\n'
    '  4. 调用硅基流动 BAAI/bge-m3 生成 1024 维向量\n'
    '  5. 以 halfvec 格式写入 knowledge_chunks 表（必须用原始 SQL）\n'
    '  6. 发布（status 2→4）\n\n'
    '为何不用 HTTP API：Windows 控制台 GBK 编码环境通过 curl 发送中文 JSON 会产生乱码。'
    'Go 程序直接操作 DB 避免了编码层问题。'
)

doc.add_page_break()

# 6. 账号密码
doc.add_heading('6. 账号密码重置汇总', level=1)

doc.add_paragraph('修复过程中因密码哈希不匹配，统一重置了所有账号的密码。')

users = [
    ('admin', 'Admin@123456', '系统管理员', '全部后台权限'),
    ('operator1', 'Opsmind@123', '运维人员', '申告处理 + 知识读写'),
    ('operator2', 'Opsmind@123', '运维人员', '申告处理 + 知识读写'),
    ('knowledge', 'Opsmind@123', '知识库管理员', '知识审核 + 发布'),
    ('admin1', 'Opsmind@123', '知识库管理员', '知识审核 + 发布'),
    ('reporter1', 'Reporter1@123', '报障人', '智能问答 + 申告提交'),
]

table4 = doc.add_table(rows=len(users)+1, cols=4, style='Light Shading Accent 1')
table4.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr4 = table4.rows[0].cells
hdr4[0].text = '用户名'
hdr4[1].text = '密码'
hdr4[2].text = '角色'
hdr4[3].text = '权限简述'
for i, (u, p, r, perm) in enumerate(users):
    row = table4.rows[i+1].cells
    row[0].text = u
    row[1].text = p
    row[2].text = r
    row[3].text = perm

doc.add_page_break()

# 7. 修改文件清单
doc.add_heading('7. 修改文件清单', level=1)

files = [
    ('后端核心修改', [
        'server/internal/service/ticket_service.go — 新增 reopen case、修复 resolve 条件、新增 writeAudit',
        'server/internal/service/user_service.go — 新增 auditRepo、writeAudit、operatorID 参数',
        'server/internal/service/role_service.go — 新增 auditRepo、writeAudit、operatorID 参数',
        'server/internal/handler/user.go — Handler 层传入 operatorID',
        'server/internal/handler/role.go — Handler 层传入 operatorID',
        'server/internal/handler/knowledge.go — 修复 ListArticles 默认 status=-1',
        'server/cmd/main.go — 注入 auditRepo 到 Service 构造函数',
    ]),
    ('前端修改', [
        'web/src/views/admin/TicketDetail.vue — 新增 reopen 按钮和样式',
        'web/src/utils/ticket.ts — actionText 新增 reopen / auto_close 映射',
    ]),
    ('新增工具', [
        'server/cmd/seed_knowledge/main.go — 知识库种子数据工具',
        'server/scripts/gen_embeddings.py — 向量嵌入生成脚本',
    ]),
    ('数据库修改', [
        'llm_configs 表 — 默认配置从 llama.cpp 改为 DeepSeek',
        'knowledge_articles 表 — 新增 6 篇运维知识文章',
        'knowledge_chunks 表 — 补全 11 个分块的 halfvec 向量嵌入',
        'users 表 — 重置所有账号密码哈希',
    ]),
]

for category, items in files:
    doc.add_heading(category, level=2)
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

# ===================== 页脚 =====================
doc.add_paragraph()
doc.add_paragraph('— 报告结束 —').alignment = WD_ALIGN_PARAGRAPH.CENTER

# 保存
output_dir = r'C:\Users\29006\Desktop'
output_path = os.path.join(output_dir, f'OpsMind_修复报告_{date.today().strftime("%Y%m%d")}.docx')
doc.save(output_path)
print(f'报告已保存：{output_path}')
